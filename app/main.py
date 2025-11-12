import streamlit as st
import streamlit.components.v1 as components
import traceback
import json
import os
from datetime import datetime
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas

# ✅ Import Chain and Portfolio
from chains import Chain
from utils import clean_text, get_page_text, is_category_url, extract_first_job_url
from portfolio import Portfolio


# ✅ Securely load Groq API key from Hugging Face Secrets
GROQ_API_KEY = st.secrets["GROQ_API_KEY"]


# ------------------ CUSTOM PURPLE PREMIUM THEME ------------------

def apply_custom_style():
    st.markdown("""
    <style>
        .stApp {
            background-color: #6B21A8 !important;
        }
        h1, h2, h3, h4, h5, h6 {
            color: #FFFFFF !important;
            font-weight: 700 !important;
        }
        p, span, div, label {
            color: #111827 !important;
        }
        .stTextInput > div > div > input {
            background-color: #FFFFFF !important;
            color: #111827 !important;
            border-radius: 8px !important;
        }
        .stAlert, .stInfo, .stSuccess, .stWarning, .stError {
            background-color: #FFFFFF !important;
            color: #111827 !important;
            border-left: 6px solid #6B21A8 !important;
            border-radius: 10px !important;
        }
        .stCode, pre, code {
            background-color: #FFFFFF !important;
            color: #111827 !important;
            border-radius: 10px !important;
            padding: 14px !important;
            overflow-x: auto !important;
            white-space: pre-wrap !important;
            word-wrap: break-word !important;
        }
        .stButton > button, .stDownloadButton > button {
            background-color: #FFFFFF !important;
            color: #6B21A8 !important;
            border-radius: 8px !important;
            padding: 10px 20px !important;
            font-weight: 600 !important;
            border: none !important;
        }
        .stButton > button:hover, .stDownloadButton > button:hover {
            background-color: #EEDBFF !important;
            color: #4A0E75 !important;
        }
    </style>
    """, unsafe_allow_html=True)


# ------------------ FILE GENERATORS ------------------

def generate_txt(subject, email_text, tone):
    return f"Subject: {subject}\nTone: {tone}\n\n{email_text}".encode("utf-8")


def generate_docx(subject, email_text, tone):
    buffer = BytesIO()
    doc = Document()
    doc.add_heading(f"Subject: {subject}", level=2)
    doc.add_paragraph(f"Tone: {tone}")
    for line in email_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf(subject, email_text, tone):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    y = 800
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(40, y, f"Subject: {subject}")
    y -= 20
    pdf.setFont("Helvetica-Oblique", 10)
    pdf.drawString(40, y, f"Tone: {tone}")
    y -= 30
    pdf.setFont("Helvetica", 11)
    for line in email_text.split("\n"):
        pdf.drawString(40, y, line)
        y -= 20
        if y < 40:
            pdf.showPage()
            y = 800
    pdf.save()
    buffer.seek(0)
    return buffer


# ------------------ COPY TO CLIPBOARD BUTTON ------------------

def copy_to_clipboard_button(email_text, key):
    safe_text = email_text.replace("`", "\\`").replace("\n", "\\n")
    button_html = f"""
        <div style="margin-top: 12px; margin-bottom: 25px; text-align: left;">
            <button id="copy_btn_{key}"
                style="
                    background-color:#FFFFFF;
                    color:#6B21A8;
                    border-radius:8px;
                    padding:8px 16px;
                    font-weight:600;
                    border:none;
                    cursor:pointer;
                    box-shadow:0px 2px 6px rgba(0,0,0,0.25);
                    transition:all 0.2s ease;">
                📋 Copy Email
            </button>
        </div>
        <script>
        const btn_{key} = document.getElementById("copy_btn_{key}");
        if (btn_{key}) {{
            btn_{key}.addEventListener("click", async () => {{
                try {{
                    await navigator.clipboard.writeText(`{safe_text}`);
                    btn_{key}.innerText = "✅ Copied!";
                    setTimeout(() => btn_{key}.innerText = "📋 Copy Email", 2000);
                }} catch (err) {{
                    btn_{key}.innerText = "❌ Failed";
                }}
            }});
        }}
        </script>
    """
    components.html(button_html, height=70)


# ------------------ EMAIL HISTORY MANAGEMENT ------------------

HISTORY_FILE = "email_history.json"

def load_history():
    if not os.path.exists(HISTORY_FILE):
        return []
    with open(HISTORY_FILE, "r", encoding="utf-8") as f:
        history = json.load(f)
    for entry in history:
        entry.setdefault("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        entry.setdefault("subject", "Generated by older version (subject unavailable)")
        entry.setdefault("tone", "Professional")
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=4, ensure_ascii=False)
    return history


def save_to_history(job_role, subject, email_text, tone):
    history = load_history()
    history.append({
        "job_role": job_role,
        "subject": subject,
        "email": email_text,
        "tone": tone,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    })
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=4, ensure_ascii=False)


def clear_history():
    if os.path.exists(HISTORY_FILE):
        os.remove(HISTORY_FILE)


# ------------------ SUBJECT LINE GENERATOR ------------------

def generate_subject_line(job):
    role = job.get("role", "position")
    company = job.get("company", "your organization")
    return f"Excited to apply for {role} at {company}"


# ------------------ STREAMLIT APP ------------------

def create_streamlit_app(chain, portfolio):
    st.set_page_config(page_title="PitchCraft AI", layout="wide")
    apply_custom_style()

    st.title("💌 PitchCraft AI — Smart Cold Email Generator")

    with st.sidebar:
        st.header("📜 Saved Email History")
        history = load_history()
        if history:
            selected = st.selectbox("View previous emails:", [f"{h['job_role']} — {h['timestamp']}" for h in history])
            if selected:
                chosen = next(h for h in history if f"{h['job_role']} — {h['timestamp']}" == selected)
                st.markdown(f"### 💼 {chosen['job_role']}")
                st.markdown(f"**📌 Subject:** {chosen['subject']}")
                st.markdown(f"**🗣️ Tone:** {chosen['tone']}")
                st.markdown(
                    f"<div style='background:#fff;border-radius:8px;padding:10px;max-height:300px;overflow:auto;'>{chosen['email']}</div>",
                    unsafe_allow_html=True
                )
                st.download_button(
                    "⬇️ Download this Email",
                    data=(f"Subject: {chosen['subject']}\nTone: {chosen['tone']}\n\n{chosen['email']}").encode("utf-8"),
                    file_name=f"{chosen['job_role']}_saved.txt"
                )
            if st.button("🗑️ Clear All History"):
                clear_history()
                st.success("History cleared successfully!")
        else:
            st.info("No saved emails yet. Generate one to save it!")

    if "jobs" not in st.session_state:
        st.session_state.jobs = []
    if "emails_generated" not in st.session_state:
        st.session_state.emails_generated = False

    url = st.text_input("Enter ANY job URL or category URL:")
    tone = st.selectbox("🎙️ Choose Email Tone:", ["Professional", "Friendly", "Confident", "Short", "Detailed"])

    if st.button("Submit"):
        try:
            url = url.strip()
            st.info("🔍 Checking URL...")
            if is_category_url(url):
                st.warning("Category URL detected — finding job listing...")
                link = extract_first_job_url(url)
                if link:
                    st.success(f"✅ Job found → {link}")
                    url = link
                else:
                    st.info("No job links found — extracting directly from page text.")
            st.info("📄 Scraping page text...")
            text = get_page_text(url)
            if not text:
                st.error("❌ Could not scrape page text.")
                return
            st.info("📁 Loading portfolio...")
            portfolio.load_portfolio()
            st.info("🧠 Extracting job information...")
            jobs = chain.extract_jobs(text)
            if not jobs:
                st.error("❌ No jobs found.")
                return
            st.session_state.jobs = jobs
            st.success(f"✅ {len(jobs)} job(s) found!")
        except Exception:
            st.error("❌ An error occurred.")
            st.code(traceback.format_exc())

    if st.session_state.jobs:
        selected_jobs = st.multiselect(
            "Select the job(s) you want to generate emails for:",
            [f"{i+1}. {job.get('role', 'Unknown Role')}" for i, job in enumerate(st.session_state.jobs)]
        )

        if st.button("✨ Generate Emails"):
            st.session_state.emails_generated = True

        if st.session_state.emails_generated:
            jobs_to_generate = [st.session_state.jobs[int(title.split('.')[0]) - 1] for title in selected_jobs]

            for i, job in enumerate(jobs_to_generate, start=1):
                with st.container():
                    st.subheader(f"💼 {job.get('role', 'Job')}")

                    job_prompt = f"[Tone: {tone}] {job}"
                    skills = job.get("skills", [])
                    links = portfolio.query_links(skills)

                    email = chain.write_mail(job_prompt, links)
                    subject = generate_subject_line(job)
                    save_to_history(job.get("role", "Unknown Role"), subject, email, tone)

                    st.markdown(f"### 📌 Subject: {subject}")
                    st.markdown(f"**🗣️ Tone:** {tone}")
                    st.markdown(
                        f"""
                        <div style="background-color:#FFFFFF;
                                   color:#111827;
                                   border-radius:10px;
                                   padding:16px;
                                   margin-top:10px;
                                   white-space:pre-wrap;
                                   word-wrap:break-word;
                                   max-height:350px;
                                   overflow-y:auto;">
                            {email}
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                    copy_to_clipboard_button(email, key=i)
                    st.download_button("📄 Download as TXT", generate_txt(subject, email, tone),
                                       f"cold_email_{i}.txt", "text/plain")
                    st.download_button("📘 Download as DOCX", generate_docx(subject, email, tone),
                                       f"cold_email_{i}.docx",
                                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    st.download_button("📕 Download as PDF", generate_pdf(subject, email, tone),
                                       f"cold_email_{i}.pdf", "application/pdf")


if __name__ == "__main__":
    chain = Chain()
    portfolio = Portfolio()
    create_streamlit_app(chain, portfolio)
